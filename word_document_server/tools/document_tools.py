"""
Document creation and manipulation tools for Word Document Server.
"""
import os
import json
import io
from typing import Dict, List, Optional, Any
from docx import Document

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension, create_document_copy
from word_document_server.utils.document_utils import get_document_properties, extract_document_text, get_document_structure, get_document_xml, insert_header_near_text, insert_line_or_paragraph_near_text
from word_document_server.core.styles import ensure_heading_style, ensure_table_style
from word_document_server.utils.azure_storage import save_document_to_storage, get_document_from_storage, get_document_url


async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document with optional metadata.

    Args:
        filename: Name of the document to create (with or without .docx extension)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    filename = ensure_docx_extension(filename)

    try:
        doc = Document()

        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author

        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)

        # Save to memory buffer first
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_data = doc_buffer.getvalue()
        doc_buffer.close()

        # Save to Azure Blob Storage (or local as fallback)
        success, message = save_document_to_storage(filename, doc_data)

        if success:
            # Try to get the public URL if available
            url = get_document_url(filename)
            if url:
                return f"Document {filename} created successfully. Access URL: {url}"
            else:
                return f"Document {filename} created successfully. {message}"
        else:
            return f"Failed to save document: {message}"

    except Exception as e:
        return f"Failed to create document: {str(e)}"


async def get_document_info(filename: str) -> str:
    """Get information about a Word document.

    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)

    try:
        # Try to get from storage first
        success, doc_data, message = get_document_from_storage(filename)

        if not success:
            return f"Document {filename} does not exist: {message}"

        # Create temporary file to analyze
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_file.write(doc_data)
            temp_path = temp_file.name

        try:
            properties = get_document_properties(temp_path)

            # Add URL if available
            url = get_document_url(filename)
            if url:
                properties['download_url'] = url

            return json.dumps(properties, indent=2)
        finally:
            os.unlink(temp_path)

    except Exception as e:
        return f"Failed to get document info: {str(e)}"


async def get_document_text(filename: str) -> str:
    """Extract all text from a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    return extract_document_text(filename)


async def get_document_outline(filename: str) -> str:
    """Get the structure of a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    structure = get_document_structure(filename)
    return json.dumps(structure, indent=2)


async def list_available_documents(directory: str = ".") -> str:
    """List all .docx files in storage.

    Args:
        directory: Directory to search for Word documents (ignored for Azure storage)
    """
    try:
        from word_document_server.utils.azure_storage import list_stored_documents

        success, file_list, message = list_stored_documents()

        if not success:
            return f"Failed to list documents: {message}"

        if not file_list:
            return "No Word documents found in storage"

        result = f"Found {len(file_list)} Word documents in storage:\n"
        for file_info in file_list:
            status = " (EXPIRED)" if file_info.get('expired', False) else ""
            size_kb = file_info.get('size', 0) / 1024
            url = get_document_url(file_info['name'])
            url_info = f"\n  URL: {url}" if url else ""
            result += f"- {file_info['name']} ({size_kb:.2f} KB){status}\n"
            result += f"  Created: {file_info.get('created', 'Unknown')}\n"
            result += f"  Expires: {file_info.get('expires', 'No expiry')}{url_info}\n"

        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"


async def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> str:
    """Create a copy of a Word document.
    
    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    source_filename = ensure_docx_extension(source_filename)
    
    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)
    
    success, message, new_path = create_document_copy(source_filename, destination_filename)
    if success:
        return message
    else:
        return f"Failed to copy document: {message}"


async def merge_documents(target_filename: str, source_filenames: List[str], add_page_breaks: bool = True) -> str:
    """Merge multiple Word documents into a single document.
    
    Args:
        target_filename: Path to the target document (will be created or overwritten)
        source_filenames: List of paths to source documents to merge
        add_page_breaks: If True, add page breaks between documents
    """
    from word_document_server.core.tables import copy_table
    
    target_filename = ensure_docx_extension(target_filename)
    
    # Check if target file is writeable
    is_writeable, error_message = check_file_writeable(target_filename)
    if not is_writeable:
        return f"Cannot create target document: {error_message}"
    
    # Validate all source documents exist
    missing_files = []
    for filename in source_filenames:
        doc_filename = ensure_docx_extension(filename)
        if not os.path.exists(doc_filename):
            missing_files.append(doc_filename)
    
    if missing_files:
        return f"Cannot merge documents. The following source files do not exist: {', '.join(missing_files)}"
    
    try:
        # Create a new document for the merged result
        target_doc = Document()
        
        # Process each source document
        for i, filename in enumerate(source_filenames):
            doc_filename = ensure_docx_extension(filename)
            source_doc = Document(doc_filename)
            
            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()
            
            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles['Normal']  # Default style
                
                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except:
                    pass
                
                # Copy run formatting
                for i, run in enumerate(paragraph.runs):
                    if i < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[i]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size
            
            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)
        
        # Save the merged document
        target_doc.save(target_filename)
        return f"Successfully merged {len(source_filenames)} documents into {target_filename}"
    except Exception as e:
        return f"Failed to merge documents: {str(e)}"


async def get_document_xml_tool(filename: str) -> str:
    """Get the raw XML structure of a Word document."""
    return get_document_xml(filename)


async def cleanup_expired_documents() -> str:
    """Clean up expired documents from storage."""
    try:
        from word_document_server.utils.azure_storage import cleanup_expired_documents

        success, count, message = cleanup_expired_documents()

        if success:
            return f"Cleanup completed: {message}"
        else:
            return f"Cleanup failed: {message}"

    except Exception as e:
        return f"Failed to cleanup documents: {str(e)}"


async def download_document(filename: str) -> str:
    """Get download URL for a document."""
    filename = ensure_docx_extension(filename)

    try:
        url = get_document_url(filename)
        if url:
            return f"Download URL for {filename}: {url}"
        else:
            return f"Document {filename} is not available for download or Azure Blob Storage is not configured"

    except Exception as e:
        return f"Failed to get download URL: {str(e)}"


async def debug_storage() -> str:
    """Debug storage configuration and available documents."""
    try:
        from word_document_server.utils.azure_storage import debug_storage_state
        return debug_storage_state()
    except Exception as e:
        return f"Failed to debug storage: {str(e)}"


async def check_document_exists(filename: str) -> str:
    """Check if a document exists in storage and provide detailed diagnostics."""
    filename = ensure_docx_extension(filename)

    try:
        # Try to get the document
        success, doc_data, message = get_document_from_storage(filename)

        result = []
        result.append(f"Document check for: {filename}")
        result.append(f"Found: {'Yes' if success else 'No'}")
        result.append(f"Message: {message}")

        if success and doc_data:
            result.append(f"Size: {len(doc_data)} bytes")

        # Also show available documents for comparison
        result.append("\n--- Available Documents ---")
        available_docs = await list_available_documents()
        result.append(available_docs)

        return "\n".join(result)

    except Exception as e:
        return f"Failed to check document existence: {str(e)}"
