"""
Template management tools for Word Document Server.
Provides MCP tools for creating, listing, and using document templates.
"""
import io
import json
from typing import Optional, Dict, Any
from docx import Document

from word_document_server.utils.file_utils import ensure_docx_extension
from word_document_server.utils.template_storage import list_templates, get_template, save_template, get_template_url
from word_document_server.utils.azure_storage import get_document_from_storage, save_document_to_storage, get_document_url


async def list_document_templates(category: str = "") -> str:
    """
    List all available document templates.

    Args:
        category: Category filter (e.g., "business", "academic", "personal"). Leave empty to list all templates.

    Returns:
        JSON string with list of templates
    """
    try:
        # Convert empty string to None for the storage function
        category_filter = category if category else None
        success, templates, message = list_templates(category_filter)

        if not success:
            return f"Failed to list templates: {message}"

        if not templates:
            category_msg = f" in category '{category}'" if category else ""
            return f"No templates found{category_msg}"

        # Format response
        result = {
            "total_templates": len(templates),
            "category_filter": category,
            "templates": []
        }

        for template in templates:
            template_info = {
                "name": template['name'],
                "category": template['category'],
                "description": template['description'],
                "author": template['author'],
                "created": template['created'],
                "size_kb": round(template['size'] / 1024, 2) if template['size'] else 0
            }

            # Add download URL if available
            url = get_template_url(template['name'], template['category'])
            if url:
                template_info['download_url'] = url

            result["templates"].append(template_info)

        return json.dumps(result, indent=2)

    except Exception as e:
        return f"Error listing templates: {str(e)}"


async def add_document_template(source_document: str, template_name: str,
                               category: str = "general", description: str = "",
                               author: str = "Unknown") -> str:
    """
    Add a document as a template to the template library.

    Automatically cleans the document by:
    - Keeping title row (row 0) and header row (row 1) in all tables, removing all data rows
    - Extracting and returning table titles for user reference
    - All other content (text, formatting, customer success info, etc.) is preserved as-is

    Args:
        source_document: Path to the source Word document to use as template
                        Can be either:
                        - A simple filename (e.g., "mydoc.docx") - will search in word-documents container
                        - A blob path (e.g., "/word-templates/Eric FER/mydoc.docx") - will use exact blob path and user folder as category
        template_name: Name for the new template (without .docx extension)
        category: Template category (default: "general", overridden by user folder if blob path provided)
        description: Description of the template
        author: Author of the template

    Returns:
        Success message with cleaning statistics, table titles list, and template URL or error description
    """
    source_document = ensure_docx_extension(source_document)

    # Clean template name (remove .docx if provided)
    if template_name.endswith('.docx'):
        template_name = template_name[:-5]

    # Track if we need to delete/replace source blob after template creation
    should_delete_source = False
    source_blob_client = None
    source_blob_name = None
    is_same_name = False  # Track if template_name matches source filename

    try:
        # Check if source_document is a blob path (starts with /word-templates/)
        if source_document.startswith('/word-templates/') or source_document.startswith('word-templates/'):
            # Remove leading slash if present
            blob_path = source_document.lstrip('/')

            # Get the document directly from the blob path
            from word_document_server.utils.azure_storage import storage

            if not storage.is_enabled():
                return "Azure Blob Storage is not configured. Cannot retrieve document from blob path."

            try:
                # Extract container name from blob path
                # blob_path is like "word-templates/Eric FER/file.docx"
                path_parts = blob_path.split('/', 1)
                if len(path_parts) < 2:
                    return f"Invalid blob path format: {source_document}. Expected format: /container-name/path/to/file.docx"

                container_name = path_parts[0]  # "word-templates"
                blob_name = path_parts[1]  # "Eric FER/file.docx"

                # Extract user folder from blob_name to use as category
                # "Eric FER/file.docx" -> "Eric FER"
                if '/' in blob_name:
                    extracted_category = blob_name.split('/')[0]
                    # Override category parameter with extracted folder name
                    category = extracted_category

                blob_client = storage.blob_service_client.get_blob_client(
                    container=container_name,
                    blob=blob_name
                )

                # Check if blob exists
                if not blob_client.exists():
                    return f"Source document not found at blob path: {source_document} (container: {container_name}, blob: {blob_name})"

                # Download the blob
                download_stream = blob_client.download_blob()
                doc_data = download_stream.readall()

                success = True
                message = f"Retrieved from blob path: {blob_path} (container: {container_name}, category: {category})"

                # Check if template_name matches source filename
                source_filename = blob_name.split('/')[-1].replace('.docx', '')
                if template_name == source_filename:
                    # Same name: save_template will overwrite the source with cleaned version (no need to delete)
                    should_delete_source = False
                    source_blob_client = None
                    source_blob_name = None
                    is_same_name = True
                else:
                    # Different name: keep source file
                    should_delete_source = False
                    source_blob_client = None

            except Exception as e:
                return f"Failed to retrieve document from blob path '{source_document}': {str(e)}"
        else:
            # Use standard document storage retrieval (word-documents container)
            success, doc_data, message = get_document_from_storage(source_document)

            # If not found in word-documents, search in word-templates container
            if not success:
                from word_document_server.utils.azure_storage import storage

                if storage.is_enabled():
                    try:
                        # Try to find the file in word-templates container
                        container_client = storage.blob_service_client.get_container_client('word-templates')

                        # List all blobs and find matching filename
                        found_blob = None
                        found_category = "general"

                        # Extract just the filename from source_document if it contains a path
                        source_filename_only = source_document.split('/')[-1]

                        for blob in container_client.list_blobs():
                            # Skip placeholder files
                            if blob.name.endswith('/.keep') or blob.name == '.keep':
                                continue

                            # Check if filename matches (case insensitive)
                            blob_filename = blob.name.split('/')[-1]  # Get just the filename
                            if blob_filename.lower() == source_filename_only.lower():
                                found_blob = blob.name
                                # Extract category from path (e.g., "Eric FER/file.docx" -> "Eric FER")
                                if '/' in blob.name:
                                    found_category = blob.name.split('/')[0]
                                break

                        if found_blob:
                            # Found the file in word-templates, download it
                            blob_client = storage.blob_service_client.get_blob_client(
                                container='word-templates',
                                blob=found_blob
                            )
                            download_stream = blob_client.download_blob()
                            doc_data = download_stream.readall()
                            success = True
                            message = f"Retrieved from word-templates: {found_blob}"
                            category = found_category  # Override category with folder name

                            # Check if template_name matches source filename
                            source_filename = found_blob.split('/')[-1].replace('.docx', '')
                            if template_name == source_filename:
                                # Same name: save_template will overwrite the source with cleaned version (no need to delete)
                                should_delete_source = False
                                source_blob_client = None
                                source_blob_name = None
                                is_same_name = True
                            else:
                                # Different name: keep source file
                                should_delete_source = False
                                source_blob_client = None
                        else:
                            return f"Source document '{source_document}' not found in word-documents or word-templates containers"
                    except Exception as e:
                        return f"Source document '{source_document}' not found in word-documents: {message}. Failed to search word-templates: {str(e)}"
                else:
                    return f"Source document '{source_document}' not found: {message}"

        # Validate it's a proper Word document by trying to open it
        try:
            doc = Document(io.BytesIO(doc_data))
            # Basic validation - check if we can access paragraphs
            _ = len(doc.paragraphs)
        except Exception as e:
            return f"Source document '{source_document}' is not a valid Word document: {str(e)}"

        # === AUTOMATIC CLEANING ===
        cleaning_stats = {
            "tables_cleaned": 0,
            "rows_removed": 0,
            "table_titles": []
        }

        # Clean all tables - keep title row (row 0) and header row (row 1)
        for table_idx, table in enumerate(doc.tables):
            # Extract title from first row (if it exists)
            table_title = ""
            if len(table.rows) > 0:
                # Get text from first row (title row)
                title_cells = []
                for cell in table.rows[0].cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        title_cells.append(cell_text)

                table_title = " - ".join(title_cells) if title_cells else f"Tableau {table_idx + 1}"
                cleaning_stats["table_titles"].append(table_title)

            # Keep first 2 rows (title + header), remove the rest
            if len(table.rows) > 2:
                rows_to_remove = len(table.rows) - 2
                # Remove rows from bottom to top to avoid index issues
                for _ in range(rows_to_remove):
                    table._element.remove(table.rows[-1]._element)
                cleaning_stats["tables_cleaned"] += 1
                cleaning_stats["rows_removed"] += rows_to_remove

        # Save the cleaned document
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        cleaned_doc_data = doc_buffer.getvalue()
        doc_buffer.close()

        # Save as template
        success, save_message = save_template(
            template_name=template_name,
            template_data=cleaned_doc_data,
            category=category,
            description=description,
            author=author
        )

        if success:
            # Handle source file status message
            deletion_info = ""
            if is_same_name:
                # Same name: file was automatically replaced by save_template with overwrite=True
                deletion_info = f"\n✅ File replaced with cleaned version (template overwrote source at same location)"
            elif should_delete_source and source_blob_client:
                # Different name but should delete source
                try:
                    source_blob_client.delete_blob()
                    deletion_info = f"\n🗑️  Source file deleted from {source_document}"
                except Exception as e:
                    deletion_info = f"\n⚠️  Warning: Could not delete source file: {str(e)}"
            else:
                # Source file kept (different name)
                if 'word-templates' in message.lower():
                    deletion_info = f"\n📁 Source file kept (different name - both source and template exist)"

            # Build cleaning summary
            cleaning_summary = ""
            if cleaning_stats['tables_cleaned'] > 0:
                cleaning_summary = f"\n✓ Template cleaned: {cleaning_stats['tables_cleaned']} table(s) processed, {cleaning_stats['rows_removed']} data row(s) removed (keeping title + header rows)"

            # Add table titles information
            tables_info = ""
            if cleaning_stats['table_titles']:
                tables_info = "\n📋 Tables in template:\n"
                for idx, title in enumerate(cleaning_stats['table_titles'], 1):
                    tables_info += f"   {idx}. {title}\n"

            # Try to get template URL
            url = get_template_url(template_name, category)
            if url:
                return f"{save_message}{cleaning_summary}{tables_info}{deletion_info}\nTemplate URL: {url}"
            else:
                return f"{save_message}{cleaning_summary}{tables_info}{deletion_info}"
        else:
            return f"Failed to save template: {save_message}"

    except Exception as e:
        return f"Error adding template: {str(e)}"


async def create_document_from_template(template_name: str, new_document_name: str,
                                      category: str = "general",
                                      variables: Dict[str, str] = {}) -> str:
    """
    Create a new document from an existing template.

    Args:
        template_name: Name of the template to use. Can be either:
                      - Just the template name (e.g., "MyTemplate") - uses category parameter
                      - Full path with category (e.g., "Eric FER/MyTemplate") - extracts category from path
        new_document_name: Name for the new document
        category: Template category (default: "general", ignored if template_name contains path)
        variables: Optional dictionary of variables to replace in the document
                  (format: {"{{variable_name}}": "replacement_value"})

    Returns:
        Success message with document URL or error description
    """
    new_document_name = ensure_docx_extension(new_document_name)

    # Clean template name (remove .docx if provided)
    if template_name.endswith('.docx'):
        template_name = template_name[:-5]

    # Check if template_name contains a path separator (category/name format)
    if '/' in template_name:
        # Extract category and actual template name from the path
        path_parts = template_name.rsplit('/', 1)  # Split from right to get last part
        if len(path_parts) == 2:
            category = path_parts[0]  # e.g., "Eric FER"
            template_name = path_parts[1]  # e.g., "Proposition commerciale Modern workplace 2024_1"

    try:
        # Get the template from storage
        success, template_data, message = get_template(template_name, category)

        if not success:
            return f"error: template not found"

        # Load the template as a Document
        doc = Document(io.BytesIO(template_data))

        # Apply variable substitutions if provided
        if variables:
            # Replace variables in paragraphs
            for paragraph in doc.paragraphs:
                for var_name, var_value in variables.items():
                    if var_name in paragraph.text:
                        # Replace in all runs to preserve formatting
                        for run in paragraph.runs:
                            if var_name in run.text:
                                run.text = run.text.replace(var_name, var_value)

            # Replace variables in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for var_name, var_value in variables.items():
                                if var_name in paragraph.text:
                                    for run in paragraph.runs:
                                        if var_name in run.text:
                                            run.text = run.text.replace(var_name, var_value)

            # Replace variables in headers and footers
            for section in doc.sections:
                # Headers
                if section.header:
                    for paragraph in section.header.paragraphs:
                        for var_name, var_value in variables.items():
                            if var_name in paragraph.text:
                                for run in paragraph.runs:
                                    if var_name in run.text:
                                        run.text = run.text.replace(var_name, var_value)

                # Footers
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        for var_name, var_value in variables.items():
                            if var_name in paragraph.text:
                                for run in paragraph.runs:
                                    if var_name in run.text:
                                        run.text = run.text.replace(var_name, var_value)

        # Save the new document
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_data = doc_buffer.getvalue()
        doc_buffer.close()

        # Save to document storage
        success, save_message = save_document_to_storage(new_document_name, doc_data)

        if success:
            # Get document URL if available
            url = get_document_url(new_document_name)

            if url:
                return f"ok\n{url}"
            else:
                return "ok"
        else:
            return f"error: {save_message}"

    except Exception as e:
        return f"error: {str(e)}"


async def get_template_info(template_name: str, category: str = "general") -> str:
    """
    Get detailed information about a specific template.

    Args:
        template_name: Name of the template. Can be either:
                      - Just the template name (e.g., "MyTemplate") - uses category parameter
                      - Full path with category (e.g., "Eric FER/MyTemplate") - extracts category from path
        category: Template category (default: "general", ignored if template_name contains path)

    Returns:
        JSON string with template information
    """
    # Clean template name (remove .docx if provided)
    if template_name.endswith('.docx'):
        template_name = template_name[:-5]

    # Check if template_name contains a path separator (category/name format)
    if '/' in template_name:
        # Extract category and actual template name from the path
        path_parts = template_name.rsplit('/', 1)  # Split from right to get last part
        if len(path_parts) == 2:
            category = path_parts[0]  # e.g., "Eric FER"
            template_name = path_parts[1]  # e.g., "Proposition commerciale Modern workplace 2024_1"

    try:
        # Get template data to analyze
        success, template_data, message = get_template(template_name, category)

        if not success:
            return f"Template '{template_name}' in category '{category}' not found: {message}"

        # Load template to analyze structure
        doc = Document(io.BytesIO(template_data))

        # Analyze template content
        info = {
            "name": template_name,
            "category": category,
            "statistics": {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "sections": len(doc.sections)
            },
            "structure": [],
            "variables_found": []
        }

        # Analyze structure and find variables
        variables_found = set()

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                para_info = {
                    "type": "paragraph",
                    "index": i,
                    "text": paragraph.text[:100] + "..." if len(paragraph.text) > 100 else paragraph.text
                }

                # Look for style information
                if paragraph.style and paragraph.style.name != 'Normal':
                    para_info["style"] = paragraph.style.name

                info["structure"].append(para_info)

                # Find variables in format {{variable_name}}
                import re
                vars_in_text = re.findall(r'\{\{([^}]+)\}\}', paragraph.text)
                variables_found.update(vars_in_text)

        # Add table information
        for i, table in enumerate(doc.tables):
            table_info = {
                "type": "table",
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0
            }
            info["structure"].append(table_info)

        info["variables_found"] = sorted(list(variables_found))

        # Add download URL if available
        url = get_template_url(template_name, category)
        if url:
            info["download_url"] = url

        return json.dumps(info, indent=2)

    except Exception as e:
        return f"Error getting template info: {str(e)}"


async def delete_document_template(template_name: str, category: str = "general") -> str:
    """
    Delete a template from the template library.

    Args:
        template_name: Name of the template to delete. Can be either:
                      - Just the template name (e.g., "MyTemplate") - uses category parameter
                      - Full path with category (e.g., "Eric FER/MyTemplate") - extracts category from path
        category: Template category (default: "general", ignored if template_name contains path)

    Returns:
        Success message or error description
    """
    # Clean template name (remove .docx if provided)
    if template_name.endswith('.docx'):
        template_name = template_name[:-5]

    # Check if template_name contains a path separator (category/name format)
    if '/' in template_name:
        # Extract category and actual template name from the path
        path_parts = template_name.rsplit('/', 1)  # Split from right to get last part
        if len(path_parts) == 2:
            category = path_parts[0]  # e.g., "Eric FER"
            template_name = path_parts[1]  # e.g., "Proposition commerciale Modern workplace 2024_1"

    try:
        from word_document_server.utils.template_storage import delete_template

        success, message = delete_template(template_name, category)
        return message

    except Exception as e:
        return f"Error deleting template: {str(e)}"