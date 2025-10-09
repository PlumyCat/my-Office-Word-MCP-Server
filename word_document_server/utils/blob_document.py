"""
Helper functions to work with documents in Azure Blob Storage.
Provides context managers and utilities for loading and saving documents.
"""
import io
from contextlib import contextmanager
from typing import Generator, Tuple
from docx import Document

from word_document_server.utils.azure_storage import get_document_from_storage, save_document_to_storage, get_document_url
from word_document_server.utils.file_utils import ensure_docx_extension


@contextmanager
def blob_document(filename: str) -> Generator[Tuple[Document, bool], None, None]:
    """
    Context manager for working with documents from Azure Blob Storage.

    Usage:
        async with blob_document('mydoc.docx') as (doc, exists):
            if not exists:
                return "Document not found"
            doc.add_paragraph("Hello")
            # Document automatically saved when context exits

    Args:
        filename: Name of the document

    Yields:
        Tuple of (Document object, exists flag)
    """
    filename = ensure_docx_extension(filename)
    doc = None
    exists = True

    # Try to get document from storage
    success, doc_data, message = get_document_from_storage(filename)

    if not success:
        exists = False
        doc = Document()  # Create new document if not found
    else:
        # Load document from blob data
        doc = Document(io.BytesIO(doc_data))

    # Yield the document for modification
    yield doc, exists

    # Save the document back to storage
    if doc is not None:
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_data = doc_buffer.getvalue()
        doc_buffer.close()

        save_document_to_storage(filename, doc_data)


def load_document_from_blob(filename: str) -> Tuple[bool, Document, str]:
    """
    Load a document from Azure Blob Storage.

    Args:
        filename: Name of the document

    Returns:
        Tuple of (success, Document object, message)
    """
    filename = ensure_docx_extension(filename)

    success, doc_data, message = get_document_from_storage(filename)

    if not success:
        return False, None, message

    try:
        doc = Document(io.BytesIO(doc_data))
        return True, doc, "Document loaded successfully"
    except Exception as e:
        return False, None, f"Failed to parse document: {str(e)}"


def save_document_to_blob(filename: str, doc: Document) -> Tuple[bool, str]:
    """
    Save a document to Azure Blob Storage.

    Args:
        filename: Name of the document
        doc: Document object to save

    Returns:
        Tuple of (success, message with optional URL)
    """
    filename = ensure_docx_extension(filename)

    try:
        # Save to memory buffer
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_data = doc_buffer.getvalue()
        doc_buffer.close()

        # Save to blob storage
        success, message = save_document_to_storage(filename, doc_data)

        if not success:
            return False, message

        # Get URL if available
        url = get_document_url(filename)
        if url:
            return True, f"Document saved. URL: {url}"
        else:
            return True, "Document saved successfully"

    except Exception as e:
        return False, f"Failed to save document: {str(e)}"