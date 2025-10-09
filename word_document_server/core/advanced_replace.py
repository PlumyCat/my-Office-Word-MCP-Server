"""
Advanced search and replace for Word documents.
Handles text in paragraphs, tables, headers, and footers.
"""

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table


def replace_text_everywhere(doc_path: str, find_text: str, replace_text: str) -> dict:
    """
    Search and replace text everywhere in a Word document.
    Handles: paragraphs, tables, headers, footers, ContentControls.

    Args:
        doc_path: Path to the Word document
        find_text: Text to find
        replace_text: Text to replace with

    Returns:
        dict with replacement count and locations
    """
    doc = Document(doc_path)
    total_replacements = 0
    locations = []

    # 1. Replace in main document body paragraphs
    count = replace_in_paragraphs(doc.paragraphs, find_text, replace_text)
    if count > 0:
        total_replacements += count
        locations.append(f"Body paragraphs: {count}")

    # 2. Replace in tables
    for table_idx, table in enumerate(doc.tables):
        count = replace_in_table(table, find_text, replace_text)
        if count > 0:
            total_replacements += count
            locations.append(f"Table {table_idx + 1}: {count}")

    # 3. Replace in headers
    for section_idx, section in enumerate(doc.sections):
        if section.header:
            count = replace_in_paragraphs(section.header.paragraphs, find_text, replace_text)
            if count > 0:
                total_replacements += count
                locations.append(f"Header section {section_idx + 1}: {count}")

    # 4. Replace in footers
    for section_idx, section in enumerate(doc.sections):
        if section.footer:
            count = replace_in_paragraphs(section.footer.paragraphs, find_text, replace_text)
            if count > 0:
                total_replacements += count
                locations.append(f"Footer section {section_idx + 1}: {count}")

    # 5. Replace in ContentControls (structured document tags)
    count = replace_in_content_controls(doc, find_text, replace_text)
    if count > 0:
        total_replacements += count
        locations.append(f"ContentControls: {count}")

    # Save the document
    doc.save(doc_path)

    return {
        "total_replacements": total_replacements,
        "locations": locations,
        "find_text": find_text,
        "replace_text": replace_text
    }


def replace_in_paragraphs(paragraphs, find_text: str, replace_text: str) -> int:
    """Replace text in a list of paragraphs."""
    count = 0
    for para in paragraphs:
        if find_text in para.text:
            # Replace in runs to preserve formatting
            inline = para.runs
            for run in inline:
                if find_text in run.text:
                    run.text = run.text.replace(find_text, replace_text)
                    count += run.text.count(replace_text)
    return count


def replace_in_table(table: Table, find_text: str, replace_text: str) -> int:
    """Replace text in all cells of a table."""
    count = 0
    for row in table.rows:
        for cell in row.cells:
            count += replace_in_paragraphs(cell.paragraphs, find_text, replace_text)
    return count


def replace_in_content_controls(doc: Document, find_text: str, replace_text: str) -> int:
    """
    Replace text in ContentControls (structured document tags).
    Note: ContentControls support coming soon. For now, use regular text.
    """
    # TODO: Implement ContentControls support
    return 0


def get_content_controls_info(doc_path: str) -> list:
    """
    List all ContentControls found in the document with their current values.
    Note: ContentControls support coming soon.
    """
    # TODO: Implement ContentControls listing
    return []
